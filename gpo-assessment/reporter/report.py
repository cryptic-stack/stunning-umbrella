from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from db import db_cursor


def _load_assessment(assessment_run_id: int):
    with db_cursor() as (_, cur):
        cur.execute(
            """
            SELECT ar.id, ar.status, ar.created_at, ar.completed_at, ps.source_type, ps.source_name,
                   COALESCE(f.name, ''), COALESCE(v.version, ''), ar.mapping_label
            FROM assessment_runs ar
            JOIN policy_sources ps ON ps.id = ar.policy_source_id
            LEFT JOIN frameworks f ON f.id = ar.framework_id
            LEFT JOIN versions v ON v.id = ar.version_id
            WHERE ar.id = %s
            """,
            (assessment_run_id,),
        )
        header = cur.fetchone()
        if not header:
            raise ValueError(f"assessment run {assessment_run_id} not found")
        cur.execute(
            """
            SELECT rule_id, setting_key, status, actual_value, expected_value, details
            FROM assessment_results
            WHERE assessment_run_id = %s
            ORDER BY id ASC
            """,
            (assessment_run_id,),
        )
        rows = cur.fetchall()
    return header, rows


def _render_markdown(summary: dict, rows: list[dict]) -> str:
    lines = [
        f"# GPO Assessment Report #{summary['assessment_run_id']}",
        "",
        f"- Source: `{summary['source_name']}` ({summary['source_type']})",
        f"- Benchmark: `{summary['framework']}` `{summary['version']}`",
        f"- Mapping: `{summary['mapping_label']}`",
        f"- Status: `{summary['status']}`",
        "",
        "## Status Counts",
        "",
    ]
    for key, value in summary["counts"].items():
        lines.append(f"- {key}: {value}")
    lines += [
        "",
        "## Results",
        "",
        "| Rule ID | Setting Key | Status | Details |",
        "|---|---|---|---|",
    ]
    for row in rows:
        lines.append(f"| {row['rule_id']} | {row['setting_key']} | {row['status']} | {row['details']} |")
    return "\n".join(lines)


def _render_html(summary: dict, rows: list[dict]) -> str:
    counts_html = "".join(f"<li><strong>{k}</strong>: {v}</li>" for k, v in summary["counts"].items())
    row_html = "".join(
        "<tr>"
        f"<td>{row['rule_id']}</td>"
        f"<td>{row['setting_key']}</td>"
        f"<td>{row['status']}</td>"
        f"<td>{row['details']}</td>"
        "</tr>"
        for row in rows
    )
    return f"""<!doctype html>
<html>
<head><meta charset="utf-8"><title>GPO Assessment {summary['assessment_run_id']}</title></head>
<body>
  <h1>GPO Assessment Report #{summary['assessment_run_id']}</h1>
  <ul>
    <li><strong>Source:</strong> {summary['source_name']} ({summary['source_type']})</li>
    <li><strong>Benchmark:</strong> {summary['framework']} {summary['version']}</li>
    <li><strong>Mapping:</strong> {summary['mapping_label']}</li>
    <li><strong>Status:</strong> {summary['status']}</li>
  </ul>
  <h2>Status Counts</h2>
  <ul>{counts_html}</ul>
  <h2>Results</h2>
  <table border="1" cellspacing="0" cellpadding="6">
    <thead><tr><th>Rule ID</th><th>Setting Key</th><th>Status</th><th>Details</th></tr></thead>
    <tbody>{row_html}</tbody>
  </table>
</body>
</html>"""


def _write_docx(path: Path, summary: dict, rows: list[dict]) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(f"GPO Assessment Report #{summary['assessment_run_id']}", level=1)
    doc.add_paragraph(f"Source: {summary['source_name']} ({summary['source_type']})")
    doc.add_paragraph(f"Benchmark: {summary['framework']} {summary['version']}")
    doc.add_paragraph(f"Mapping: {summary['mapping_label']}")
    doc.add_paragraph(f"Status: {summary['status']}")

    doc.add_heading("Status Counts", level=2)
    for key, value in summary["counts"].items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Results", level=2)
    table = doc.add_table(rows=1, cols=4)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Rule ID"
    header_cells[1].text = "Setting Key"
    header_cells[2].text = "Status"
    header_cells[3].text = "Details"
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["rule_id"]
        cells[1].text = row["setting_key"]
        cells[2].text = row["status"]
        cells[3].text = row["details"]
    doc.save(path.as_posix())


def export_assessment_report(assessment_run_id: int, output_dir: str) -> dict:
    import xlsxwriter

    header, raw_rows = _load_assessment(assessment_run_id)
    rows = [
        {
            "rule_id": row[0] or "",
            "setting_key": row[1] or "",
            "status": row[2] or "",
            "actual_value": row[3] or {},
            "expected_value": row[4] or {},
            "details": row[5] or "",
        }
        for row in raw_rows
    ]
    counts = Counter(row["status"] for row in rows)
    summary = {
        "assessment_run_id": header[0],
        "status": header[1],
        "created_at": str(header[2]),
        "completed_at": str(header[3]) if header[3] else "",
        "source_type": header[4] or "",
        "source_name": header[5] or "",
        "framework": header[6] or "",
        "version": header[7] or "",
        "mapping_label": header[8] or "",
        "counts": dict(counts),
    }

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    base = f"gpo_assessment_{assessment_run_id}"

    md_path = out / f"{base}.md"
    html_path = out / f"{base}.html"
    csv_path = out / f"{base}.csv"
    xlsx_path = out / f"{base}.xlsx"
    docx_path = out / f"{base}.docx"
    json_path = out / f"{base}.json"

    md_path.write_text(_render_markdown(summary, rows), encoding="utf-8")
    html_path.write_text(_render_html(summary, rows), encoding="utf-8")
    json_path.write_text(json.dumps({"summary": summary, "results": rows}, indent=2), encoding="utf-8")

    with csv_path.open("w", encoding="utf-8", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(["rule_id", "setting_key", "status", "details", "actual_value", "expected_value"])
        for row in rows:
            writer.writerow(
                [
                    row["rule_id"],
                    row["setting_key"],
                    row["status"],
                    row["details"],
                    json.dumps(row["actual_value"]),
                    json.dumps(row["expected_value"]),
                ]
            )

    workbook = xlsxwriter.Workbook(xlsx_path.as_posix())
    summary_ws = workbook.add_worksheet("summary")
    summary_ws.write_row(0, 0, ["field", "value"])
    for idx, (key, value) in enumerate(summary.items(), start=1):
        summary_ws.write_row(idx, 0, [key, json.dumps(value) if isinstance(value, dict) else str(value)])

    results_ws = workbook.add_worksheet("results")
    results_ws.write_row(0, 0, ["rule_id", "setting_key", "status", "details", "actual_value", "expected_value"])
    for idx, row in enumerate(rows, start=1):
        results_ws.write_row(
            idx,
            0,
            [
                row["rule_id"],
                row["setting_key"],
                row["status"],
                row["details"],
                json.dumps(row["actual_value"]),
                json.dumps(row["expected_value"]),
            ],
        )
    workbook.close()

    _write_docx(docx_path, summary, rows)

    return {
        "json": str(json_path),
        "markdown": str(md_path),
        "html": str(html_path),
        "csv": str(csv_path),
        "xlsx": str(xlsx_path),
        "docx": str(docx_path),
    }
